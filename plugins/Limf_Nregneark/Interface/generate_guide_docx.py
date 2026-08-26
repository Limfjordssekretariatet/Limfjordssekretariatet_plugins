"""
Genererer Interface Guide som Word-dokument (.docx).
Kør med: python generate_guide_docx.py
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color_hex):
    """Sæt baggrundsfarve på en tabelcelle."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shading_elm)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Tilføj en formateret tabel med header-farve."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '1565C0')

    # Rækker
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'E3F2FD')

    # Kolonnebredder
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # lidt luft efter tabellen


def add_note(doc, text, prefix="BEMÆRK"):
    """Tilføj en fremhævet note-paragraf."""
    p = doc.add_paragraph()
    run_prefix = p.add_run(f"  {prefix}: ")
    run_prefix.bold = True
    run_prefix.font.size = Pt(10)
    run_prefix.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    run_text = p.add_run(text)
    run_text.font.size = Pt(10)
    run_text.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_tip(doc, text):
    add_note(doc, text, prefix="TIP")


def add_warning(doc, text):
    add_note(doc, text, prefix="ADVARSEL")


def add_important(doc, text):
    add_note(doc, text, prefix="VIGTIGT")


def main():
    doc = Document()

    # Sæt standard font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ================================================================
    # TITEL
    # ================================================================
    title = doc.add_heading('Guide til Modelmappe Interfacet', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph(
        'Interfacet er et QGIS-plugin (PyQt-dialog) til at opsætte projekter, '
        'hente data fra Dataforsyningen og køre geoprocesserings-modeller. '
        'Hele dialogen er scrollbar og opdelt i 5 sektioner (bokse).'
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in intro.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # ================================================================
    # OVERORDNET STRUKTUR
    # ================================================================
    doc.add_heading('Overordnet struktur', level=1)

    add_styled_table(doc,
        ['Nr.', 'Sektion', 'Formål'],
        [
            ['1', 'Boks 1 — Opsætning', 'Vælg mappe, projektnavn, scripts, projektområde og token'],
            ['—', 'Ekstra projektområder', 'Tilføj flere projektområder og navngiv dem'],
            ['2', 'Boks 2 — Forberedende data',
             'Beregn strømningsveje, oplande og klip grunddata'],
            ['3', 'Boks 3 — Tilførsel', 'Vandopland, Direkte opland, Projektområde (planlagt)'],
            ['4', 'Boks 4 — Omsætning', 'Oversvømmelse, Overrisling, Ekstensivering (planlagt)'],
            ['—', 'Close-knap', 'Lukker interfacet'],
        ],
        col_widths=[1.5, 5, 10]
    )

    # Låselogik
    doc.add_heading('Låselogik', level=2)
    p = doc.add_paragraph(
        'Boks 2, 3 og 4 er låst (grå og utilgængelige) indtil alle tre felter i Boks 1 er udfyldt:'
    )
    doc.add_paragraph('Mappe er valgt', style='List Bullet')
    doc.add_paragraph('Projektområde-fil er valgt', style='List Bullet')
    doc.add_paragraph('Token er indtastet', style='List Bullet')
    doc.add_paragraph(
        'Så snart alle tre har tekst, låses boks 2–4 automatisk op.'
    )

    # ================================================================
    # BOKS 1
    # ================================================================
    doc.add_heading('Boks 1 — Opsætning', level=1)
    doc.add_paragraph(
        'Denne boks samler al den grundlæggende input, som resten af interfacet afhænger af.'
    )

    # --- Udpeg mappe ---
    doc.add_heading('1. Udpeg vedlagte mappe', level=2)
    add_styled_table(doc,
        ['Element', 'Beskrivelse'],
        [
            ['Tekstfelt (lineEditMappe)',
             'Skrivebeskyttet felt der viser stien til den valgte mappe. Placeholder: "Ingen mappe valgt..."'],
            ['"Vælg mappe"-knap (btnVaelgMappe)',
             'Åbner en standard mappevalgs-dialog. Den valgte mappe bruges som rodmappe for alt output: modeller, scripts og outputfiler.'],
            ['✕-knap (btnRydMappe)',
             'Rød nulstil-knap. Rydder det valgte mappefelt. Vises kun når feltet har indhold.'],
            ['?-ikon (hjaelpMappe)',
             'Blåt hjælpe-badge. Hover over det for vejledning (tooltip).'],
        ],
        col_widths=[5, 12]
    )
    add_important(doc,
        'Denne mappe er den vedlagte projektmappe (Vaadomraade_Modeller-mappen) der indeholder '
        'undermapperne Modeller/, Scripts/, Grunddata/ osv. Det er IKKE en outputmappe du opretter — '
        'det er selve hovedmappen med alle filer.'
    )

    # --- Navngiv projekt ---
    doc.add_heading('2. Navngiv projekt', level=2)
    add_styled_table(doc,
        ['Element', 'Beskrivelse'],
        [
            ['Tekstfelt (lineEditProjektnavn)',
             'Her skriver du et navn til dit projekt, fx "Projekt_Gudenå".'],
            ['✓-knap grøn (btnOpretProjekt)',
             'Opretter en projektmappe med det angivne navn under den valgte mappe. '
             'Der oprettes også en .qgz-projektfil inde i mappen (hvis den ikke allerede findes). '
             'Knappen vises først når feltet har tekst.'],
            ['?-ikon (hjaelpProjektnavn)',
             'Tooltip: "Skriv et projektnavn og tryk på det grønne flueben for at oprette '
             'projektmappen og en .qgz-projektfil under den valgte mappe."'],
        ],
        col_widths=[5, 12]
    )
    doc.add_paragraph('Hvad sker der når man trykker ✓?', style='List Bullet')
    doc.add_paragraph('Opretter mappen: <valgt mappe>/<projektnavn>/', style='List Bullet 2')
    doc.add_paragraph('Gemmer en QGIS-projektfil: <valgt mappe>/<projektnavn>/<projektnavn>.qgz', style='List Bullet 2')
    doc.add_paragraph('Hvis projektfilen allerede findes, overskriver den IKKE.', style='List Bullet 2')

    # --- Projektområde ---
    doc.add_heading('3. Tilføj dit projektområde (.shp)', level=2)
    add_styled_table(doc,
        ['Element', 'Beskrivelse'],
        [
            ['Tekstfelt (lineEditProjektomraade)',
             'Skrivebeskyttet felt der viser stien til den valgte projektområde-fil. '
             'Placeholder: "Ingen fil valgt..."'],
            ['"Vælg fil"-knap (btnVaelgFil)',
             'Åbner en fildialog filtreret til geodata-formater: .shp, .gpkg, .geojson, .kml, .gml.'],
            ['✕-knap (btnRydProjektomraade)',
             'Rød nulstil-knap. Rydder feltet. Vises kun når feltet har indhold.'],
            ['?-ikon (hjaelpProjektomraade)',
             'Hjælpe-badge.'],
        ],
        col_widths=[5, 12]
    )
    add_important(doc,
        'Denne fil er det primære projektområde (shapefil) der bruges som input til alle modeller. '
        'Det er typisk en polygon der afgrænser det geografiske område, du arbejder med.'
    )

    # --- Token ---
    doc.add_heading('5. Indsæt Dataforsyningen token', level=2)
    add_styled_table(doc,
        ['Element', 'Beskrivelse'],
        [
            ['Tekstfelt (lineEditToken)',
             'Her indtastes dit personlige Dataforsyningen-token (API-nøgle). '
             'Placeholder: "Indtast token..."'],
            ['?-ikon (hjaelpToken)',
             'Hjælpe-badge.'],
        ],
        col_widths=[5, 12]
    )
    add_note(doc,
        'Tokenet bruges når terrænet skal hentes fra Dataforsyningens API. '
        'Du kan oprette et token på dataforsyningen.dk.'
    )

    # ================================================================
    # EKSTRA PROJEKTOMRÅDER
    # ================================================================
    doc.add_heading('Ekstra projektområder', level=1)
    doc.add_paragraph(
        'Denne sektion sidder i sin egen boks mellem Boks 1 og Boks 2.'
    )

    add_styled_table(doc,
        ['Element', 'Beskrivelse'],
        [
            ['Titel', '"Tilføj ét eller flere projektområder:" (fed tekst)'],
            ['Projektområde 2+ felt + "Vælg fil"-knap',
             'Vælg en ekstra projektområde-fil. Feltet er skrivebeskyttet — brug knappen.'],
            ['✕-knap', 'Ryd det pågældende ekstra projektområde-felt. Vises kun når feltet har indhold.'],
            ['Auto-tilføj', 'Når du udfylder det seneste ekstra felt, tilføjes automatisk et nyt (op til 10 ekstra).'],
            ['?-ikon (hjaelpEkstraProjekter)', 'Hjælpe-badge.'],
        ],
        col_widths=[5, 12]
    )

    doc.add_heading('Navngiv projektområde', level=2)
    add_styled_table(doc,
        ['Element', 'Beskrivelse'],
        [
            ['Tekstfelt (lineEditProjektomraadeNavn)',
             'Skriv et navn for projektområdet, fx "Opland_Nord".'],
            ['✓-knap grøn (btnOpretProjektomraadeMappe)',
             'Opretter en undermappe: <valgt mappe>/<projektnavn>/<projektområdenavn>/. '
             'Kræver at projektmappen allerede er oprettet via ✓ ovenfor.'],
            ['?-ikon (hjaelpProjektomraadeNavn)',
             'Tooltip: "Skriv et navn for projektområdet og tryk på det grønne flueben..."'],
        ],
        col_widths=[5, 12]
    )
    add_warning(doc,
        'Projektmappen SKAL eksistere først! Hvis du ikke har trykket ✓ ved '
        '"Navngiv projekt", får du en advarsel.'
    )

    # ================================================================
    # BOKS 2
    # ================================================================
    doc.add_heading('Boks 2 — Forberedende data', level=1)
    doc.add_paragraph('Låst indtil Boks 1 er komplet.')
    doc.add_paragraph(
        'Denne boks indeholder modeller der forbereder de grundlæggende geodata.'
    )

    # Højdemodel
    doc.add_heading('Højdemodel', level=2)
    add_styled_table(doc,
        ['Element', 'Beskrivelse'],
        [
            ['Label', '"Højdemodel"'],
            ['"Beregn strømningsveje" (btnKorHojdemodel)',
             'Kører "Beregn strømningsveje". Findes der et præberegnet grundlag for området, hentes det — ellers hentes terrænet via Dataforsyningen '
             '(bruger dit token) og genererer to outputfiler i <mappe>/Outputfiler/:'],
            ['', '• Hoejdemodel.tif — det konditionerede terræn'],
            ['', '• Channel_Networks.gpkg — vandløbsnetværk'],
            ['?-ikon (hjaelpHojdemodel)', 'Hjælpe-badge.'],
        ],
        col_widths=[5, 12]
    )
    add_note(doc,
        'Under kørsel vises en progressbar. Knappen ændrer sig til "Kører..." og låses. '
        'Du kan annullere via dialogen.'
    )

    # Oplande
    doc.add_heading('Udvælg oplande model', level=2)
    add_styled_table(doc,
        ['Element', 'Beskrivelse'],
        [
            ['Label', '"Udvælg oplande model"'],
            ['"Kør"-knap (btnKorOplande)',
             'Kører "Beregn oplande". Bruger projektområdet som input. '
             'Denne model bruger de outputfiler (Højdemodel, Channel_Networks) der blev skabt '
             'i trinnet ovenfor.'],
            ['?-ikon (hjaelpOplande)', 'Hjælpe-badge.'],
        ],
        col_widths=[5, 12]
    )
    add_important(doc,
        'Kør altid "Beregn strømningsveje" FØR Udvælg oplande — ellers mangler '
        'inputfilerne.'
    )

    # Klip grunddata
    doc.add_heading('Klip grunddata model', level=2)
    add_styled_table(doc,
        ['Element', 'Beskrivelse'],
        [
            ['Label', '"Klip grunddata model"'],
            ['"Kør"-knap (btnKorKlipGrunddata)',
             'Kører Klip_Grunddata_Model.model3. Klipper grunddata til dit projektområde '
             'og gemmer følgende i <mappe>/Outputfiler/:'],
            ['', '• Jordbund_Klippet.gpkg'],
            ['', '• Marker_klippet.gpkg'],
            ['', '• Befaestet_Areal.gpkg'],
            ['', '• Natur.gpkg'],
            ['', '• Omdriftsarealer.gpkg'],
            ['?-ikon (hjaelpKlipGrunddata)', 'Hjælpe-badge.'],
        ],
        col_widths=[5, 12]
    )

    # ================================================================
    # BOKS 3
    # ================================================================
    doc.add_heading('Boks 3 — Tilførsel', level=1)
    doc.add_paragraph('Låst indtil Boks 1 er komplet.')
    doc.add_paragraph(
        'Denne boks omhandler tilførsel af vand/stof. '
        'Bemærk: Ingen af disse har Kør-knapper implementeret endnu — de er forberedt i layoutet.'
    )

    add_styled_table(doc,
        ['Element', 'Beskrivelse'],
        [
            ['Vandopland', 'Label + hjælpe-ikon. Forventes at tilføje en vandoplande-model.'],
            ['Direkte opland', 'Label + hjælpe-ikon. Forventes at tilføje en direkte-opland-model.'],
            ['Projektområde', 'Label + hjælpe-ikon. En ekstra projektområde-kontekst for tilførsels-beregninger.'],
        ],
        col_widths=[5, 12]
    )
    add_note(doc, 'Boks 3 har endnu ingen Kør-knapper — elementerne er layoutet men afventer implementation.')

    # ================================================================
    # BOKS 4
    # ================================================================
    doc.add_heading('Boks 4 — Omsætning', level=1)
    doc.add_paragraph('Låst indtil Boks 1 er komplet.')
    doc.add_paragraph(
        'Denne boks omhandler omsætnings-modeller. Ligesom Boks 3 er dette delvist forberedt.'
    )

    add_styled_table(doc,
        ['Element', 'Beskrivelse'],
        [
            ['(Oversvømmelse)', 'Label i parentes — markerer at modellen endnu ikke er implementeret.'],
            ['Overrisling model', 'Label + hjælpe-ikon. Forventes at modellere overrisling.'],
            ['(Ekstensivering)', 'Label i parentes — markerer at modellen endnu ikke er implementeret.'],
        ],
        col_widths=[5, 12]
    )
    add_note(doc, 'Elementer i (parenteser) er planlagte men endnu ikke implementerede modeller.')

    # ================================================================
    # CLOSE-KNAP
    # ================================================================
    doc.add_heading('Close-knap', level=1)
    doc.add_paragraph(
        'Nederst i dialogen (udenfor scroll-området) sidder en standard Close-knap '
        'der lukker hele interfacet.'
    )

    # ================================================================
    # ANBEFALET WORKFLOW
    # ================================================================
    doc.add_heading('Anbefalet arbejdsgang (workflow)', level=1)

    workflow_steps = [
        ('1.', 'Vælg mappe', 'Peg på den vedlagte Vaadomraade_Modeller-mappe'),
        ('2.', 'Navngiv projekt', 'Skriv projektnavnet og tryk ✓ for at oprette mapper'),
        ('3.', 'Vælg projektområde', 'Vælg din .shp-fil'),
        ('4.', 'Indtast token', 'Din Dataforsyningen API-nøgle'),
        ('5.', 'Beregn strømningsveje', 'Skaffer terrænet og udleder strømningen'),
        ('6.', 'Kør Udvælg oplande', 'Beregner oplande baseret på højdemodellen'),
        ('7.', 'Kør Klip grunddata', 'Klipper grunddata til dit projektområde'),
    ]

    add_styled_table(doc,
        ['Trin', 'Handling', 'Beskrivelse'],
        [(s, h, b) for s, h, b in workflow_steps],
        col_widths=[1.5, 4.5, 11]
    )

    # ================================================================
    # IKONER
    # ================================================================
    doc.add_heading('Ikoner og visuelle elementer', level=1)

    add_styled_table(doc,
        ['Ikon', 'Betydning'],
        [
            ['? (blå cirkel)', 'Hjælpe-badge — hover for tooltip-vejledning'],
            ['✕ (rød)', 'Nulstil/ryd felt — vises kun når feltet har indhold'],
            ['✓ (grøn)', 'Bekræft/opret — opretter mapper eller filer'],
            ['Kør', 'Starter den pågældende geoprocesseringsmodel'],
        ],
        col_widths=[4, 13]
    )

    # ================================================================
    # MAPPESTRUKTUR
    # ================================================================
    doc.add_heading('Mappestruktur der oprettes', level=1)
    doc.add_paragraph(
        'Følgende viser den samlede mappestruktur som interfacet arbejder med og opretter:'
    )

    struktur = (
        '<valgt mappe>/\n'
        '├── Grunddata/\n'
        '├── Interface/\n'
        '├── Modeller/\n'
        '│   └── Klip_Grunddata_Model.model3\n'
        '├── Outputfiler/\n'
        '│   ├── Hoejdemodel.tif\n'
        '│   ├── Channel_Networks.gpkg\n'
        '│   ├── Jordbund_Klippet.gpkg\n'
        '│   ├── Marker_klippet.gpkg\n'
        '│   ├── Befaestet_Areal.gpkg\n'
        '│   ├── Natur.gpkg\n'
        '│   └── Omdriftsarealer.gpkg\n'
        '├── Projekthej/\n'
        '├── Scripts/\n'
        '└── <projektnavn>/\n'
        '    ├── <projektnavn>.qgz\n'
        '    └── <projektområdenavn>/'
    )

    p = doc.add_paragraph()
    run = p.add_run(struktur)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    # ================================================================
    # GEM
    # ================================================================
    output_path = os.path.join(os.path.dirname(__file__), 'Interface_Guide.docx')
    doc.save(output_path)
    print(f'Word-fil gemt: {output_path}')


if __name__ == '__main__':
    main()
